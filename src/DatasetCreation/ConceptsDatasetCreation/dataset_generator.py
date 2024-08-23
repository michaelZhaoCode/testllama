import json
import os
import random
import re
from datetime import datetime

import pandas as pd
import yaml
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from openai import OpenAI

import src.helper_functions as helper
from src.DatasetCreation.ConceptsDatasetCreation.constants import (
    estimated_total_num_examples,
)
from src.DatasetCreation.ConceptsDatasetCreation.constants import generation_model_name
from src.DatasetCreation.ConceptsDatasetCreation.constants import openai_api_key
from src.DatasetCreation.ConceptsDatasetCreation.constants import openai_org_id
from src.DatasetCreation.ConceptsDatasetCreation.OpenAI_base_class import openai_base
from src.DatasetCreation.ConceptsDatasetCreation.OpenAI_pricing import (
    openai_pricing_dict,
)


# Set the seed
random.seed(1)


# single_label_generator class is used to create examples for a single label of a dataset
class single_label_generator(openai_base):
    def __init__(
        self,
        investigated_concept,
        generated_num_examples,
        investigator_name,
        prompt_version,
    ):
        super().__init__(
            generation_model_name,
            investigated_concept,
            investigator_name,
            prompt_version,
        )

        self.generated_num_examples = generated_num_examples
        self.prompt_settings = [
            "Workplace",
            "Academia",
            "Sports",
            "Entrepreneurship",
            "Politics",
            "Arts",
            "Music",
            "Community",
            "Science",
            "Technology",
            "Social Media",
        ]

        self.conversation = {}
        self.generated_dataset_str = None
        self.generated_dataset_list = None

    @property
    def session_id(self):
        return f"LabelGenerator-{id(self)}"

    def adjust_num_examples(self, prompt):
        return re.sub(r"{num_examples}", str(self.generated_num_examples), prompt)

    def load_prompt(self, investigated_concept, prompt_version):
        prompt_file_name = f"{investigated_concept}_{prompt_version}.yaml"
        file_path = os.path.join(
            os.getcwd(),
            "src",
            "DatasetCreation",
            "ConceptsDatasetCreation",
            "prompts",
            "dataset_generation",
            self.investigator_name,
            prompt_file_name,
        )
        # file_path = f"src/ConceptsDatasetCreation/prompts/dataset_generation/{concept}_prompt_{version}.yaml"
        assert os.path.exists(file_path), (
            f"generation prompt version {prompt_version} for concept {investigated_concept} "
            f"does not exist \n file path requested: {file_path}"
        )
        with open(file_path, encoding="utf-8") as file:
            data = yaml.safe_load(file)
        prompt = data["text"]

        adjusted_prompt = self.adjust_num_examples(prompt)

        return adjusted_prompt

    def predict(self, openai_obj, assistant_input=""):
        # use assistant input to add already generated examples, so that they exist in GPT4's context
        prompt = self.load_prompt(self.investigated_concept, self.prompt_version)
        prediction, usage = self.make_OpenAI_API_call(
            openai_obj, prompt, assistant_input
        )

        # calculate API call cost
        self.generation_prompt_tokens = usage.prompt_tokens
        self.generation_completion_tokens = usage.completion_tokens
        API_usage_cost = self.calculate_API_cost(
            self.model, usage.prompt_tokens, usage.completion_tokens
        )
        self.generation_usage = API_usage_cost
        self.estimated_whole_dataset_generation_cost = (
            self.calculate_estimated_total_API_cost(
                self.model, usage.prompt_tokens, usage.completion_tokens
            )
        )

        self.generated_dataset_str = prediction
        self.generated_dataset_list = self.convert_prediction_list(prediction)

        # get average number of words in the generated examples
        avg_num_words = helper.get_avg_num_words(self.generated_dataset_list)
        print(
            f"average number of words for {self.investigated_concept}: {avg_num_words:.1f}"
        )

        # save dataset to json and csv files
        self.save_prediction_json(self.generated_dataset_list)
        self.save_prediction_csv(self.generated_dataset_list)

        return prediction, API_usage_cost

    def convert_prediction_list(self, prediction, note: str = None):
        # first, check that the generated text is given in a numbered list with the number of examples equal to
        #           the required number
        # '\d+\.\s': \d+ means a single/multi-numbered integer, \. means a '.', \s means a space or tab
        # possible regex pattern: r'\d+\..+?\n'
        re_pattern = r"\d+\.\s|\d+:\s|\d+-\s"
        if len(re.findall(re_pattern, prediction)) < self.generated_num_examples:
            self.save_prediction_docx(prediction, note)
            raise Exception(
                "generated number of examples is either not equal to required number "
                "or isn't formatted in a numbered list"
            )
        else:
            # remove any text before the numbered list
            match = re.search(re_pattern, prediction)
            numbered_list_str = prediction[match.start() :]
            # remove new lines from string
            numbered_list_str = re.sub(r"\n", "", numbered_list_str)
            # Split the string on the pattern of digits followed by a period and a space
            generated_examples = re.split(re_pattern, numbered_list_str)
            # The first element will be an empty string if the string starts with a number, so we remove it
            generated_examples = [example for example in generated_examples if example]

            return generated_examples

    def save_prediction_docx(self, prediction, file_name_note: str = None):
        # Create a new Document object
        doc = Document()

        # Add the string input as a paragraph
        doc.add_paragraph(prediction)

        # Save the document to a .docx file
        save_dir = self.set_dataset_save_dir(file_extension="docx")
        os.makedirs(save_dir, exist_ok=True)  # create directory if it doesn't exist

        file_name = self.create_dataset_file_name(save_dir)

        if file_name_note is not None:
            file_name = file_name + "_" + file_name_note

        file_path = os.path.join(save_dir, f"{file_name}.docx")
        doc.save(file_path)

        print("prediction saved to output.docx")

        return

    def save_all_docx(self, file_name_note: str = "conversation"):
        def set_col_widths(table):
            widths = [Inches(3), Inches(5)]
            for row in table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width

        # Create a new Document object
        doc = Document()

        doc.add_heading("All Prompts and Generated examples", level=1)

        section = doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        for concept, generation in self.conversation.items():
            # add table with all contents for the concept
            table = doc.add_table(rows=3, cols=2)

            table.style = "Table Grid"
            hdr_row_1, hdr_row_2 = table.rows[0].cells
            hdr_row = hdr_row_1.merge(hdr_row_2)
            hdr_row.text = "Concept: " + concept

            hdr_row_1, hdr_row_2 = table.rows[1].cells
            prompt_row = hdr_row_1.merge(hdr_row_2)
            prompt_row.text = generation[0]["content"]

            hdr_cells = table.rows[2].cells
            hdr_cells[0].text = "Continuation Prompts"
            hdr_cells[1].text = "Generated Examples"

            set_col_widths(table)

            for i in range(2, len(generation), 2):
                entry = generation[i - 1]
                output = generation[i]
                prompt = entry["content"]
                content = output["content"]
                row_cells = table.add_row().cells
                row_cells[0].text = prompt
                row_cells[1].text = content
                set_col_widths(table)

            doc.add_page_break()

        # Save the document to a .docx file
        save_dir = self.set_dataset_save_dir(file_extension="docx")
        os.makedirs(save_dir, exist_ok=True)  # create directory if it doesn't exist

        file_name = self.create_dataset_file_name(save_dir)

        if file_name_note is not None:
            file_name = file_name + "_" + file_name_note

        file_path = os.path.join(save_dir, f"{file_name}.docx")
        doc.save(file_path)

        print(f"All prompts and predictions saved to {file_path}")

        return

    def save_prediction_json(
        self,
        generated_examples_list,
        file_name_note: str = None,
        call_single=False,
        investigated_concept=None,
    ):
        # first convert list to dictionary
        generated_examples_dict = {
            i + 1: value for i, value in enumerate(generated_examples_list)
        }

        # set the save directory
        if call_single is True:
            save_dir = single_label_generator.set_dataset_save_dir(
                self, file_extension="json", investigated_concept=investigated_concept
            )
        else:
            save_dir = self.set_dataset_save_dir(
                file_extension="json", investigated_concept=investigated_concept
            )
        os.makedirs(save_dir, exist_ok=True)  # create directory if it doesn't exist

        file_name = self.create_dataset_file_name(save_dir)

        if file_name_note is not None:
            file_name = file_name + "_" + file_name_note

        file_path = os.path.join(save_dir, f"{file_name}.json")

        # Open a file for writing and use json.dump to write the dictionary to the file
        with open(file_path, "w") as json_file:
            json.dump(generated_examples_dict, json_file, indent=4)

        return

    def save_prediction_csv(self, generated_examples_list, file_name_note: str = None):
        # set the save directory
        save_dir = self.set_dataset_save_dir(file_extension="csv")
        os.makedirs(save_dir, exist_ok=True)  # create directory if it doesn't exist

        file_name = self.create_dataset_file_name(save_dir)

        if file_name_note is not None:
            file_name = file_name + "_" + file_name_note

        file_path = os.path.join(save_dir, f"{file_name}.csv")
        print(f"saved file in {file_path}")

        # Convert to DataFrame
        df = pd.DataFrame(generated_examples_list, columns=["input examples"])

        # Save to CSV
        df.to_csv(file_path, index=False)

        return

    def set_dataset_save_dir(self, file_extension, investigated_concept=None):
        if investigated_concept:
            return os.path.join(
                os.getcwd(),
                "data",
                "concept_probing",
                "generation",
                "single_label",
                investigated_concept,
                self.investigator_name,
                file_extension,
            )
        else:
            return os.path.join(
                os.getcwd(),
                "data",
                "concept_probing",
                "generation",
                "single_label",
                self.investigated_concept,
                self.investigator_name,
                file_extension,
            )

    def create_dataset_file_name(self, save_dir):
        # Get the current date and time
        date = str(datetime.now().date())

        # investigator_name = self.investigator_name
        prompt_version = f"P{self.prompt_version}"

        # Get the list of all entries in the directory
        entries = os.listdir(save_dir)

        # regex to get the version number without including the extension
        # regex_pattern = r'([^_]+)(?=\.[^.]+$)'
        regex_pattern = r"_(\d+)"
        # get the dataset version for each dataset in the directory
        dataset_versions = []
        for entry in entries:
            if "DS_Store" not in entry:
                dataset_versions.append(int(re.search(regex_pattern, entry).group(1)))
        # dataset_versions = [int(re.search(regex_pattern, i).group(1)) for i in entries]
        dataset_versions.sort()
        # print(f"dataset versions: {dataset_versions}")
        # get new dataset version number
        new_dataset_version = (
            0 if not dataset_versions else int(dataset_versions[-1]) + 1
        )

        file_name = f"{date}_{prompt_version}_{new_dataset_version}"

        return file_name

    def calculate_estimated_total_API_cost(
        self, openai_model, prompt_tokens, completion_tokens
    ):
        estimated_total_prompt_tokens = prompt_tokens
        estimated_total_completion_tokens = completion_tokens * int(
            self.estimated_total_num_examples / self.generated_num_examples
        )

        return self.calculate_API_cost(
            openai_model,
            estimated_total_prompt_tokens,
            estimated_total_completion_tokens,
        )


class single_label_generator_chat(single_label_generator):
    def __init__(
        self,
        investigated_concept,
        generated_num_examples,
        incremental_num_examples,
        investigator_name,
        prompt_version,
        example_strings_file_name,
    ):
        super().__init__(
            investigated_concept,
            generated_num_examples,
            investigator_name,
            prompt_version,
        )

        self.generated_num_examples = generated_num_examples

        assert (
            self.generated_num_examples % incremental_num_examples == 0
        ), "generated number of examples must be divisible by incremental number of examples"
        self.incremental_num_examples = incremental_num_examples
        self.example_strings = self.load_example_strings(example_strings_file_name)

    def load_example_strings(self, example_strings_file_name):
        if ".txt" not in example_strings_file_name:
            example_strings_file_name = f"{example_strings_file_name}.txt"
        file_path = os.path.join(
            os.getcwd(),
            "data",
            "concept_probing",
            "generation",
            "example_sentences",
            example_strings_file_name,
        )

        assert os.path.exists(file_path), "example strings file name is not found"

        with open(file_path) as file:
            examples = [line.strip() for line in file]

        # shuffle list
        # random.shuffle(examples)

        return examples

    def load_prompt(self, investigated_concept):
        prompt_file_name = f"{investigated_concept}_{self.prompt_version}.yaml"
        file_path = os.path.join(
            os.getcwd(),
            "src",
            "DatasetCreation",
            "ConceptsDatasetCreation",
            "prompts",
            "dataset_generation",
            self.investigator_name,
            prompt_file_name,
        )
        assert os.path.exists(file_path), (
            f"generation prompt version {self.prompt_version} for concept {investigated_concept} "
            f"does not exist \n file path requested: {file_path}"
        )
        with open(file_path, encoding="utf-8") as file:
            data = yaml.safe_load(file)
        prompt = data["text"]
        continuation_prompt = data["continuation_text"]

        adjusted_continuation_prompt = self.adjust_num_examples(continuation_prompt)

        return prompt, adjusted_continuation_prompt

    def adjust_num_examples(self, prompt):
        return re.sub(r"{num_examples}", str(self.incremental_num_examples), prompt)

    def add_examples_to_continuation_prompt(self, continuation_prompt, sublist):
        adjusted_continuation_prompt = continuation_prompt
        for i, example in enumerate(sublist):
            adjusted_continuation_prompt += f"{i+1}: {example}\n"

        return adjusted_continuation_prompt

    def save_prediction_csv(self, generated_examples_list, file_name_note: str = None):
        # set the save directory
        save_dir = self.set_dataset_save_dir(file_extension="csv")
        os.makedirs(save_dir, exist_ok=True)  # create directory if it doesn't exist

        file_name = self.create_dataset_file_name(save_dir)

        if file_name_note is not None:
            file_name = file_name + "_" + file_name_note

        file_path = os.path.join(save_dir, f"{file_name}.csv")
        print(f"saved file in {file_path}")

        # Convert to DataFrame
        df = pd.DataFrame(
            generated_examples_list, columns=["input examples", "GPT4 generation label"]
        )

        # Save to CSV
        df.to_csv(file_path, index=False)

        return

    def make_OpenAI_Chatbot_calls(
        self, openai_obj, prompt, continuation_prompt, investigated_concept
    ):
        number_conversation_turns = int(
            self.generated_num_examples / self.incremental_num_examples
        )
        msgs = [{"role": "system", "name": "generator", "content": prompt}]
        prediction_all = ""
        prompt_tokens_all = 0
        completion_tokens_all = 0
        for i in range(number_conversation_turns):
            # get examples strings and add them to the continuation prompt
            example_strings_sublist = self.example_strings[
                self.incremental_num_examples
                * i : self.incremental_num_examples
                * (i + 1)
            ]
            adjusted_continuation_prompt = self.add_examples_to_continuation_prompt(
                continuation_prompt, example_strings_sublist
            )

            if "{context}" in adjusted_continuation_prompt:
                context = self.prompt_settings[i % 11]
                adjusted_continuation_prompt = re.sub(
                    r"{context}", str(context), adjusted_continuation_prompt
                )

            msgs.append(
                {
                    "role": "user",
                    "name": "director",
                    "content": adjusted_continuation_prompt,
                }
            )
            prediction, usage = self.create_chat_message(
                openai_obj=openai_obj, messages=msgs, model=self.model
            )
            msgs.append(
                {"role": "assistant", "name": "generator", "content": prediction}
            )
            prediction_all = prediction_all + "\n" + prediction
            prompt_tokens_all += usage.prompt_tokens
            completion_tokens_all += usage.completion_tokens

            # save predictions to docx file
            current_num_generations = (i + 1) * self.incremental_num_examples
            if current_num_generations % 200 == 0:
                self.save_prediction_docx(prediction_all, f"{investigated_concept}")

            print(f"generated {self.incremental_num_examples*(i+1)} examples")

        if investigated_concept in self.conversation:
            self.conversation[investigated_concept].append(msgs)
        else:
            self.conversation[investigated_concept] = msgs
        return prediction_all, prompt_tokens_all, completion_tokens_all

    def make_OpenAI_Chatbot_calls_v1(
        self, openai_obj, prompt, continuation_prompt, investigated_concept
    ):
        number_conversation_turns = int(
            self.generated_num_examples / self.incremental_num_examples
        )
        prediction_all = ""
        prompt_tokens_all = 0
        completion_tokens_all = 0
        for i in range(number_conversation_turns):
            # get examples strings and add them to the continuation prompt
            example_strings_sublist = self.example_strings[
                self.incremental_num_examples
                * i : self.incremental_num_examples
                * (i + 1)
            ]
            adjusted_continuation_prompt = self.add_examples_to_continuation_prompt(
                continuation_prompt, example_strings_sublist
            )

            if "{context}" in adjusted_continuation_prompt:
                context = self.prompt_settings[i % 11]
                adjusted_continuation_prompt = re.sub(
                    r"{context}", str(context), adjusted_continuation_prompt
                )

            msgs = [{"role": "system", "name": "generator", "content": prompt}]
            msgs.append(
                {
                    "role": "user",
                    "name": "director",
                    "content": adjusted_continuation_prompt,
                }
            )
            prediction, usage = self.create_chat_message(
                openai_obj=openai_obj, messages=msgs, model=self.model
            )
            prediction_all = prediction_all + "\n" + prediction
            prompt_tokens_all += usage.prompt_tokens
            completion_tokens_all += usage.completion_tokens

            # save predictions to docx file
            current_num_generations = (i + 1) * self.incremental_num_examples
            if current_num_generations % 200 == 0:
                self.save_prediction_docx(prediction_all, f"{investigated_concept}")

            print(f"generated {self.incremental_num_examples*(i+1)} examples")

        if investigated_concept in self.conversation:
            self.conversation[investigated_concept].append(msgs)
        else:
            self.conversation[investigated_concept] = msgs
        return prediction_all, prompt_tokens_all, completion_tokens_all

    def generate_class_examples(self, openai_obj, investigated_concept):
        prompt, continuation_prompt = self.load_prompt(investigated_concept)
        (
            predicted_examples,
            prompt_tokens,
            completion_tokens,
        ) = self.make_OpenAI_Chatbot_calls_v1(
            openai_obj, prompt, continuation_prompt, investigated_concept
        )

        predicted_examples_list = self.convert_prediction_list(
            predicted_examples, note=f"{investigated_concept}"
        )

        # # save dataset to json file
        # self.save_prediction_json(
        #     predicted_examples_list,
        #     call_single=True,
        #     investigated_concept=investigated_concept,
        # )

        # get average number of words in the positive examples
        avg_num_words = helper.get_avg_num_words(predicted_examples_list)
        print(
            f"average number of words for {investigated_concept}: {avg_num_words:.1f}"
        )

        self.generation_prompt_tokens += prompt_tokens
        self.generation_completion_tokens += completion_tokens

        return predicted_examples_list

    def predict(self, openai_obj):
        # generate +1 examples
        examples = self.generate_class_examples(
            openai_obj, f"{self.investigated_concept}"
        )

        self.save_all_docx()

        # calculate API call cost
        API_usage_cost = self.calculate_API_cost(
            self.model, self.generation_prompt_tokens, self.generation_completion_tokens
        )
        self.generation_usage = API_usage_cost
        self.estimated_whole_dataset_generation_cost = (
            self.calculate_estimated_total_API_cost(
                self.model,
                self.generation_prompt_tokens,
                self.generation_completion_tokens,
            )
        )

        """ combine datasets and randomize """
        whole_dataset = examples
        # add labels to the dataset
        if "not_" in self.investigated_concept:
            labels = [0] * len(examples)
        elif "lack_" in self.investigated_concept:
            labels = [-1] * len(examples)
        else:
            labels = [1] * len(examples)

        labelled_dataset = [[a, b] for a, b in zip(whole_dataset, labels)]

        self.generated_dataset_str = f"{examples}"
        self.generated_dataset_list = labelled_dataset

        # save dataset to json and csv files
        self.save_prediction_json(self.generated_dataset_list)
        # self.save_prediction_csv(self.generated_dataset_list, file_name_note="labelled")
        self.save_prediction_csv(self.generated_dataset_list)

        return self.generated_dataset_str, API_usage_cost

    def calculate_estimated_total_API_cost(
        self, openai_model, prompt_tokens, completion_tokens
    ):
        # Step 1: Get the system prompt tokens - Multiply by num_examples/incremental_num_examples
        # Load system & continuation prompt
        main_prompt, continuation_prompt = self.load_prompt(self.investigated_concept)
        message = [{"role": "system", "content": main_prompt}]

        # Retrieve input and output price
        if openai_model in openai_pricing_dict:
            input_price, output_price = openai_pricing_dict[openai_model]
        else:
            # Maybe use an exception here?
            print(f"ERROR: MODEL UNKNOWN ({openai_model})")

        # Create a new instance of chatbot
        openai_obj = OpenAI(organization=openai_org_id, api_key=openai_api_key)

        # Get system bot input tokens (TODO: The testing part also needs a price update)
        response, usage = self.create_chat_message(
            openai_obj, message, generation_model_name
        )
        sys_prompt_input_tokens = usage.prompt_tokens

        # Get number of generations needed for 1/3 of the entire dataset (3 categories)
        num_generations = (
            estimated_total_num_examples / self.incremental_num_examples
        ) / 3

        # Get estimated system prompt tokens (1/3)
        sys_prompt_estimated_tokens = sys_prompt_input_tokens * num_generations

        # Get estimated system prompt price (1/3)
        sys_prompt_estimation_price = input_price * sys_prompt_estimated_tokens
        print(f"sys_prompt price is approx: {sys_prompt_estimation_price}")

        # Step 2: Get the continuation prompt - 1 + 2 + 3 + ... + num_examples/incremental_num_examples times
        new_message = [{"role": "system", "content": continuation_prompt}]
        new_openai_obj = OpenAI(organization=openai_org_id, api_key=openai_api_key)
        new_response, new_usage = self.create_chat_message(
            new_openai_obj, new_message, generation_model_name
        )

        # Get continuation prompt input tokens (TODO: The testing part also needs a price update)
        continuation_prompt_input_tokens = new_usage.prompt_tokens

        # Get number of time continuation prompt is generated (1/3)
        num_times_cont_prompt = helper.arithmatic_series(
            1, num_generations, num_generations
        )

        # Get estimated continuation system prompt tokens (1/3)
        continuation_prompt_estimated_tokens = (
            continuation_prompt_input_tokens * num_times_cont_prompt
        )

        # Estimate continuation prompt price (1/3)
        cont_prompt_estimation_price = (
            input_price * continuation_prompt_estimated_tokens
        )

        # Estimate continuation prompt price (entire dataset)
        print(f"cont_prompt price is approx: {cont_prompt_estimation_price}")

        # Step 3: Get the generation tokens - Get the total number of completion tokens, divide them by the
        #         number of examples generated, then scale it up.
        # Get the approximate # of tokens for each example (total output tokens/total number of examples)
        avg_example_output_tokens = completion_tokens / (
            self.generated_num_examples * 3
        )

        # Get the output price estimation (entire dataset)
        output_price_estimation = (
            estimated_total_num_examples * avg_example_output_tokens * output_price
        )
        print(f"output price is approx: {output_price_estimation}")

        # Subtract total input tokens from the system prompt and continuation prompt tokens, divided by
        #       number of sentences needed in context, then scale it up (1/3 of the data)
        # Get the number of generations needed to create the current dataset (not the estimated one)
        num_generations_curr_session = (
            self.generated_num_examples
        ) / self.incremental_num_examples

        # Tokens left for context in this session = total input tokens - sys_prompt_tokens in current session -
        #           continuation prompt tokens in current session
        context_input_tokens = (
            prompt_tokens
            - 3 * sys_prompt_input_tokens * num_generations_curr_session
            - 3
            * helper.arithmatic_series(
                1, num_generations_curr_session, num_generations_curr_session
            )
            * continuation_prompt_input_tokens
        )

        # Total number of examples in the context
        num_context_elements = helper.arithmatic_series(
            self.incremental_num_examples * 2,
            num_generations_curr_session * self.incremental_num_examples * 2,
            num_generations_curr_session,
        )

        # Get the average example context tokens
        avg_context_example_input_token = context_input_tokens / num_context_elements

        # Get context price estimation (1/3): input price * number of examples in the context for the larger
        #           dataset (1/3) * average example context tokens
        context_price_estimation = (
            input_price
            * helper.arithmatic_series(
                self.incremental_num_examples * 2,
                num_generations * self.incremental_num_examples * 2,
                num_generations,
            )
            * avg_context_example_input_token
        )

        # Get context price estimation (entire dataset)
        print(f"context price is approx: {context_price_estimation}")

        total_price = (
            sys_prompt_estimation_price
            + cont_prompt_estimation_price
            + output_price_estimation
            + context_price_estimation
        )
        print(f"total price is approx: {total_price}")
        # 40k example price
        print(
            f"40k example price is approx: {total_price*40000/estimated_total_num_examples*3}"
        )
        return 0


class whole_dataset_generator(single_label_generator):
    def __init__(
        self,
        investigated_concept,
        generated_num_examples,
        investigator_name,
        prompt_version,
    ):
        super().__init__(
            investigated_concept,
            generated_num_examples,
            investigator_name,
            prompt_version,
        )

    def predict(self, openai_obj, pos_assistant_input="", neg_assistant_input=""):
        # use positive and negative assistant inputs to add already generated examples, so that they exist
        #           in GPT4's context

        # generate +ve examples
        pos_prompt = self.load_prompt(self.investigated_concept, self.prompt_version)
        pos_prediction, pos_usage = self.make_OpenAI_API_call(
            openai_obj, pos_prompt, pos_assistant_input
        )
        pos_examples = self.convert_prediction_list(pos_prediction, note="pos")
        # get average number of words in the positive examples
        avg_num_words = helper.get_avg_num_words(pos_examples)
        print(
            f"average number of words for {self.investigated_concept} +ve: {avg_num_words:.1f}"
        )

        # generate -ve examples
        neg_prompt = self.load_prompt(
            f"not_{self.investigated_concept}", self.prompt_version
        )
        neg_prediction, neg_usage = self.make_OpenAI_API_call(
            openai_obj, neg_prompt, neg_assistant_input
        )
        neg_examples = self.convert_prediction_list(neg_prediction, note="neg")
        # get average number of words in the positive examples
        avg_num_words = helper.get_avg_num_words(neg_examples)
        print(
            f"average number of words for {self.investigated_concept} -ve: {avg_num_words:.1f}"
        )

        # calculate API call cost
        self.generation_prompt_tokens += (
            pos_usage.prompt_tokens + neg_usage.prompt_tokens
        )
        self.generation_completion_tokens += (
            pos_usage.completion_tokens + neg_usage.completion_tokens
        )
        API_usage_cost = self.calculate_API_cost(
            self.model, self.generation_prompt_tokens, self.generation_completion_tokens
        )
        self.generation_usage = API_usage_cost
        self.estimated_whole_dataset_generation_cost = (
            self.calculate_estimated_total_API_cost(
                self.model,
                self.generation_prompt_tokens,
                self.generation_completion_tokens,
            )
        )

        """ combine datasets and randomize """
        whole_dataset = pos_examples + neg_examples
        # add labels to the dataset
        labels = [1] * len(pos_examples) + [0] * len(neg_examples)
        labelled_dataset = [[a, b] for a, b in zip(whole_dataset, labels)]
        # Shuffle the dataset to mix between +ve and -ve examples
        random.shuffle(labelled_dataset)

        self.generated_dataset_str = f"{pos_prediction} \n\n\n{neg_prediction}"
        self.generated_dataset_list = labelled_dataset

        # save dataset to json and csv files
        self.save_prediction_json(self.generated_dataset_list)
        # self.save_prediction_csv(self.generated_dataset_list, file_name_note="labelled")
        self.save_prediction_csv(self.generated_dataset_list)

        return self.generated_dataset_str, API_usage_cost

    def set_dataset_save_dir(self, file_extension, investigated_concept=None):
        if investigated_concept:
            return os.path.join(
                os.getcwd(),
                "data",
                "concept_probing",
                "generation",
                "whole_dataset",
                investigated_concept,
                self.investigator_name,
                file_extension,
            )
        else:
            return os.path.join(
                os.getcwd(),
                "data",
                "concept_probing",
                "generation",
                "whole_dataset",
                self.investigated_concept,
                self.investigator_name,
                file_extension,
            )

    def calculate_estimated_total_API_cost(
        self, openai_model, prompt_tokens, completion_tokens
    ):
        estimated_total_prompt_tokens = prompt_tokens
        # we multiply by 2 because "self.generated_num_examples" represents number of examples for a single class
        estimated_total_completion_tokens = completion_tokens * int(
            self.estimated_total_num_examples / (self.generated_num_examples * 2)
        )

        return self.calculate_API_cost(
            openai_model,
            estimated_total_prompt_tokens,
            estimated_total_completion_tokens,
        )

    def save_prediction_csv(self, generated_examples_list, file_name_note: str = None):
        # set the save directory
        save_dir = self.set_dataset_save_dir(file_extension="csv")
        os.makedirs(save_dir, exist_ok=True)  # create directory if it doesn't exist

        file_name = self.create_dataset_file_name(save_dir)

        if file_name_note is not None:
            file_name = file_name + "_" + file_name_note

        file_path = os.path.join(save_dir, f"{file_name}.csv")
        print(f"saved file in {file_path}")

        # Convert to DataFrame
        df = pd.DataFrame(
            generated_examples_list, columns=["input examples", "GPT4 generation label"]
        )

        # Save to CSV
        df.to_csv(file_path, index=False)

        return


class whole_dataset_generator_chat(whole_dataset_generator):
    def __init__(
        self,
        investigated_concept,
        generated_num_examples,
        incremental_num_examples,
        investigator_name,
        prompt_version,
        example_strings_file_name,
    ):
        super().__init__(
            investigated_concept,
            generated_num_examples,
            investigator_name,
            prompt_version,
        )

        assert (
            generated_num_examples % 3 == 0
        ), "generated number of examples for the whole dataset must be divisible by 3"
        # this attribute internally still represents the number of examples generated for a single class,
        #       so it has to be divided by 3
        self.generated_num_examples = int(generated_num_examples / 3)

        assert (
            self.generated_num_examples % incremental_num_examples == 0
        ), "generated number of examples must be divisible by incremental number of examples"
        self.incremental_num_examples = incremental_num_examples

        self.example_strings = self.load_example_strings(example_strings_file_name)

    def load_example_strings(self, example_strings_file_name):
        if ".txt" not in example_strings_file_name:
            example_strings_file_name = f"{example_strings_file_name}.txt"
        file_path = os.path.join(
            os.getcwd(),
            "data",
            "concept_probing",
            "generation",
            "example_sentences",
            example_strings_file_name,
        )

        assert os.path.exists(file_path), "example strings file name is not found"

        with open(file_path) as file:
            examples = [line.strip() for line in file]

        # shuffle list
        random.shuffle(examples)

        return examples

    def load_prompt(self, investigated_concept):
        prompt_file_name = f"{investigated_concept}_{self.prompt_version}.yaml"
        file_path = os.path.join(
            os.getcwd(),
            "src",
            "DatasetCreation",
            "ConceptsDatasetCreation",
            "prompts",
            "dataset_generation",
            self.investigator_name,
            prompt_file_name,
        )
        assert os.path.exists(file_path), (
            f"generation prompt version {self.prompt_version} for concept {investigated_concept} "
            f"does not exist \n file path requested: {file_path}"
        )
        with open(file_path, encoding="utf-8") as file:
            data = yaml.safe_load(file)
        prompt = data["text"]
        continuation_prompt = data["continuation_text"]

        adjusted_continuation_prompt = self.adjust_num_examples(continuation_prompt)

        return prompt, adjusted_continuation_prompt

    def adjust_num_examples(self, prompt):
        return re.sub(r"{num_examples}", str(self.incremental_num_examples), prompt)

    def add_examples_to_continuation_prompt(self, continuation_prompt, sublist):
        adjusted_continuation_prompt = continuation_prompt
        for i, example in enumerate(sublist):
            adjusted_continuation_prompt += f"{i+1}: {example}\n"

        return adjusted_continuation_prompt

    def make_OpenAI_Chatbot_calls(
        self, openai_obj, prompt, continuation_prompt, investigated_concept
    ):
        number_conversation_turns = int(
            self.generated_num_examples / self.incremental_num_examples
        )
        msgs = [{"role": "system", "name": "generator", "content": prompt}]
        prediction_all = ""
        prompt_tokens_all = 0
        completion_tokens_all = 0
        for i in range(number_conversation_turns):
            # get examples strings and add them to the continuation prompt
            example_strings_sublist = self.example_strings[
                self.incremental_num_examples
                * i : self.incremental_num_examples
                * (i + 1)
            ]
            adjusted_continuation_prompt = self.add_examples_to_continuation_prompt(
                continuation_prompt, example_strings_sublist
            )

            if "{context}" in adjusted_continuation_prompt:
                context = self.prompt_settings[i % 11]
                adjusted_continuation_prompt = re.sub(
                    r"{context}", str(context), adjusted_continuation_prompt
                )

            print(adjusted_continuation_prompt)

            msgs.append(
                {
                    "role": "user",
                    "name": "director",
                    "content": adjusted_continuation_prompt,
                }
            )
            prediction, usage = self.create_chat_message(
                openai_obj=openai_obj, messages=msgs, model=self.model
            )
            print(prediction)
            msgs.append(
                {"role": "assistant", "name": "generator", "content": prediction}
            )
            prediction_all = prediction_all + "\n" + prediction
            prompt_tokens_all += usage.prompt_tokens
            completion_tokens_all += usage.completion_tokens

            # save predictions to docx file
            current_num_generations = (i + 1) * self.incremental_num_examples
            if current_num_generations % 1000 == 0:
                self.save_prediction_docx(prediction_all, f"{investigated_concept}")

        if investigated_concept in self.conversation:
            self.conversation[investigated_concept].append(msgs)
        else:
            self.conversation[investigated_concept] = msgs

        return prediction_all, prompt_tokens_all, completion_tokens_all

    def generate_class_examples(self, openai_obj, investigated_concept):
        prompt, continuation_prompt = self.load_prompt(investigated_concept)
        (
            predicted_examples,
            prompt_tokens,
            completion_tokens,
        ) = self.make_OpenAI_Chatbot_calls(
            openai_obj, prompt, continuation_prompt, investigated_concept
        )

        predicted_examples_list = self.convert_prediction_list(
            predicted_examples, note=f"{investigated_concept}"
        )

        # save dataset to json file
        self.save_prediction_json(
            predicted_examples_list,
            call_single=True,
            investigated_concept=investigated_concept,
        )

        # get average number of words in the positive examples
        avg_num_words = helper.get_avg_num_words(predicted_examples_list)
        print(
            f"average number of words for {investigated_concept}: {avg_num_words:.1f}"
        )

        self.generation_prompt_tokens += prompt_tokens
        self.generation_completion_tokens += completion_tokens

        return predicted_examples_list

    def predict(self, openai_obj):
        # generate +1 examples
        pos_examples = self.generate_class_examples(
            openai_obj, f"{self.investigated_concept}"
        )

        # generate -1 examples
        neg_examples = self.generate_class_examples(
            openai_obj, f"lack_{self.investigated_concept}"
        )
        # generate 0 examples
        null_examples = self.generate_class_examples(
            openai_obj, f"not_{self.investigated_concept}"
        )

        self.save_all_docx()

        # calculate API call cost
        API_usage_cost = self.calculate_API_cost(
            self.model, self.generation_prompt_tokens, self.generation_completion_tokens
        )
        self.generation_usage = API_usage_cost
        self.estimated_whole_dataset_generation_cost = (
            self.calculate_estimated_total_API_cost(
                self.model,
                self.generation_prompt_tokens,
                self.generation_completion_tokens,
            )
        )

        """ combine datasets and randomize """
        whole_dataset = pos_examples + null_examples + neg_examples
        # add labels to the dataset
        labels = (
            [1] * len(pos_examples)
            + [0] * len(null_examples)
            + [-1] * len(neg_examples)
        )
        labelled_dataset = [[a, b] for a, b in zip(whole_dataset, labels)]
        # Shuffle the dataset to mix between +1, -1, and 0 examples
        random.shuffle(labelled_dataset)

        self.generated_dataset_str = (
            f"{pos_examples} \n\n\n{null_examples} \n\n\n{neg_examples}"
        )
        self.generated_dataset_list = labelled_dataset

        # save dataset to json and csv files
        self.save_prediction_json(self.generated_dataset_list)
        # self.save_prediction_csv(self.generated_dataset_list, file_name_note="labelled")
        self.save_prediction_csv(self.generated_dataset_list)

        return self.generated_dataset_str, API_usage_cost

    def calculate_estimated_total_API_cost(
        self, openai_model, prompt_tokens, completion_tokens
    ):
        # Step 1: Get the system prompt tokens - Multiply by num_examples/incremental_num_examples
        # Load system & continuation prompt
        main_prompt, continuation_prompt = self.load_prompt(self.investigated_concept)
        message = [{"role": "system", "content": main_prompt}]

        # Retrieve input and output price
        if openai_model in openai_pricing_dict:
            input_price, output_price = openai_pricing_dict[openai_model]
        else:
            # Maybe use an exception here?
            print(f"ERROR: MODEL UNKNOWN ({openai_model})")

        # Create a new instance of chatbot
        openai_obj = OpenAI(organization=openai_org_id, api_key=openai_api_key)

        # Get system bot input tokens (TODO: The testing part also needs a price update)
        response, usage = self.create_chat_message(
            openai_obj, message, generation_model_name
        )
        sys_prompt_input_tokens = usage.prompt_tokens

        # Get number of generations needed for 1/3 of the entire dataset (3 categories)
        num_generations = (
            estimated_total_num_examples / self.incremental_num_examples
        ) / 3

        # Get estimated system prompt tokens (1/3)
        sys_prompt_estimated_tokens = sys_prompt_input_tokens * num_generations

        # Get estimated system prompt price (1/3)
        sys_prompt_estimation_price = input_price * sys_prompt_estimated_tokens

        # Get estimated system prompt price (entire dataset)
        sys_prompt_estimation_price *= 3
        print(f"sys_prompt price is approx: {sys_prompt_estimation_price}")

        # Step 2: Get the continuation prompt - 1 + 2 + 3 + ... + num_examples/incremental_num_examples times
        new_message = [{"role": "system", "content": continuation_prompt}]
        new_openai_obj = OpenAI(organization=openai_org_id, api_key=openai_api_key)
        new_response, new_usage = self.create_chat_message(
            new_openai_obj, new_message, generation_model_name
        )

        # Get continuation prompt input tokens (TODO: The testing part also needs a price update)
        continuation_prompt_input_tokens = new_usage.prompt_tokens

        # Get number of time continuation prompt is generated (1/3)
        num_times_cont_prompt = helper.arithmatic_series(
            1, num_generations, num_generations
        )

        # Get estimated continuation system prompt tokens (1/3)
        continuation_prompt_estimated_tokens = (
            continuation_prompt_input_tokens * num_times_cont_prompt
        )

        # Estimate continuation prompt price (1/3)
        cont_prompt_estimation_price = (
            input_price * continuation_prompt_estimated_tokens
        )

        # Estimate continuation prompt price (entire dataset)
        cont_prompt_estimation_price *= 3
        print(f"cont_prompt price is approx: {cont_prompt_estimation_price}")

        # Step 3: Get the generation tokens - Get the total number of completion tokens, divide them by the
        #           number of examples generated, then scale it up.
        # Get the approximate # of tokens for each example (total output tokens/total number of examples)
        avg_example_output_tokens = completion_tokens / (
            self.generated_num_examples * 3
        )

        # Get the output price estimation (entire dataset)
        output_price_estimation = (
            estimated_total_num_examples * avg_example_output_tokens * output_price
        )
        print(f"output price is approx: {output_price_estimation}")

        # Subtract total input tokens from the system prompt and continuation prompt tokens, divided by
        #           number of sentences needed in context, then scale it up (1/3 of the data)
        # Get the number of generations needed to create the current dataset (not the estimated one)
        num_generations_curr_session = (
            self.generated_num_examples
        ) / self.incremental_num_examples

        # Tokens left for context in this session = total input tokens - sys_prompt_tokens in current session -
        #           continuation prompt tokens in current session
        context_input_tokens = (
            prompt_tokens
            - 3 * sys_prompt_input_tokens * num_generations_curr_session
            - 3
            * helper.arithmatic_series(
                1, num_generations_curr_session, num_generations_curr_session
            )
            * continuation_prompt_input_tokens
        )

        # Total number of examples in the context
        num_context_elements = helper.arithmatic_series(
            self.incremental_num_examples * 2,
            num_generations_curr_session * self.incremental_num_examples * 2,
            num_generations_curr_session,
        )

        # Get the average example context tokens
        avg_context_example_input_token = context_input_tokens / num_context_elements

        # Get context price estimation (1/3): input price * number of examples in the context for the larger
        #           dataset (1/3) * average example context tokens
        context_price_estimation = (
            input_price
            * helper.arithmatic_series(
                self.incremental_num_examples * 2,
                num_generations * self.incremental_num_examples * 2,
                num_generations,
            )
            * avg_context_example_input_token
        )

        # Get context price estimation (entire dataset)
        context_price_estimation *= 3
        print(f"context price is approx: {context_price_estimation}")

        total_price = (
            sys_prompt_estimation_price
            + cont_prompt_estimation_price
            + output_price_estimation
            + context_price_estimation
        )
        print(f"total price is approx: {total_price}")
        # 40k example price
        print(
            f"40k example price is approx: {total_price*40000/estimated_total_num_examples}"
        )
        return 0
